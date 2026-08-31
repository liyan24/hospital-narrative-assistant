# -*- coding: utf-8 -*-
"""将 docs/科研助手说明文档.md 原样转换为 Word 文档（一次性工具）。"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
MD = ROOT / "docs" / "科研助手说明文档.md"
OUT = ROOT / "docs" / "科研助手说明文档.docx"
BASE = MD.parent  # 图片相对路径基准


def style_run(run, font="宋体", size=11, bold=False, italic=False,
              color=None, mono=False):
    run.font.name = "Consolas" if mono else "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color:
        run.font.color.rgb = RGBColor(*color)


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_rich_text(p, text, size=11, font="宋体", base_bold=False):
    """处理行内 **加粗** 与 `代码`"""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            style_run(p.add_run(part[2:-2]), font=font, size=size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            style_run(p.add_run(part[1:-1]), size=size - 0.5, mono=True)
        else:
            style_run(p.add_run(part), font=font, size=size, bold=base_bold)


def add_para(doc, text, size=11, font="宋体", bold=False, indent=False,
             align=None, space_after=6, color=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.35
    if indent:
        pf.first_line_indent = Pt(size * 2)
    if align:
        p.alignment = align
    if color:
        style_run(p.add_run(text), font=font, size=size, bold=bold, color=color)
    else:
        add_rich_text(p, text, size=size, font=font, base_bold=bold)
    return p


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    t = doc.add_table(rows=len(rows), cols=len(header))
    t.style = "Table Grid"
    for j, cell in enumerate(header):
        c = t.rows[0].cells[j]
        c.paragraphs[0].text = ""
        add_rich_text(c.paragraphs[0], cell, size=10, base_bold=True)
    for i, row in enumerate(body, start=1):
        for j, cell in enumerate(row):
            c = t.rows[i].cells[j]
            c.paragraphs[0].text = ""
            add_rich_text(c.paragraphs[0], cell, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image(doc, rel_path, alt):
    fp = (BASE / rel_path).resolve()
    if not fp.exists():
        add_para(doc, f"[图片缺失：{rel_path}]", size=9, color=(200, 0, 0))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 页面截图较宽，报告页较窄，按宽高比缩放
    from PIL import Image  # noqa: 若无可退回固定宽度
    try:
        w, h = Image.open(fp).size
        width_cm = 15.5 if w / h > 1.2 else 12.5
    except Exception:
        width_cm = 15.5
    p.add_run().add_picture(str(fp), width=Cm(width_cm))
    if alt:
        add_para(doc, alt, size=9, indent=False,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10,
                 color=(90, 90, 90))


def flush_table(doc, buf):
    rows = [[c.strip() for c in line.strip().strip("|").split("|")]
            for line in buf]
    # 去掉分隔行（--- 行）
    rows = [r for r in rows if not all(set(c) <= set(":- ") for c in r)]
    if rows:
        add_table(doc, rows)


def main():
    doc = Document()
    lines = MD.read_text(encoding="utf-8").splitlines()
    i = 0
    table_buf = []
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                style_run(p.add_run("\n".join(code_buf)), size=8.5, mono=True)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # 表格
        if line.strip().startswith("|"):
            table_buf.append(line)
            i += 1
            continue
        elif table_buf:
            flush_table(doc, table_buf)
            table_buf = []

        stripped = line.strip()

        # 标题
        if stripped.startswith("# "):
            add_para(doc, stripped[2:], size=18, font="黑体", bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
        elif stripped.startswith("## "):
            add_para(doc, stripped[3:], size=14, font="黑体", bold=True,
                     space_after=8)
        elif stripped.startswith("### "):
            add_para(doc, stripped[4:], size=12, font="黑体", bold=True,
                     space_after=6)
        # 图片
        elif stripped.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
            if m:
                add_image(doc, m.group(2), m.group(1))
        # 引用块
        elif stripped.startswith(">"):
            add_para(doc, stripped.lstrip("> "), size=9.5, color=(90, 90, 90))
        # 分割线
        elif stripped == "---":
            pass
        # 列表
        elif stripped.startswith("- "):
            add_para(doc, "• " + stripped[2:], space_after=3)
        elif re.match(r"^\d+\.\s", stripped):
            add_para(doc, stripped, space_after=3)
        # 普通段落
        elif stripped:
            add_para(doc, stripped)

        i += 1

    if table_buf:
        flush_table(doc, table_buf)

    doc.save(str(OUT))
    print("saved:", OUT)


if __name__ == "__main__":
    main()
