"""自动科研流水线端到端验证（真实 LLM 调用，勿启动 uvicorn）。
用法：.venv/Scripts/python.exe -X utf8 scripts/verify_auto_research.py
"""
import json
import sys
import time
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from services.research.auto_research_service import auto_research_service
from services.research.skills.registry import SKILL_REGISTRY

PAPER_SECTIONS = ["摘要", "前言", "资料与方法", "结果", "讨论", "结论"]
TIMEOUT = 600  # 10 分钟


def step(title: str):
    print(f"\n{'=' * 60}\n{title}\n{'=' * 60}", flush=True)


def main():
    # ========== ① 议题推荐 ==========
    step("① propose_topics() 真实调用")
    topics = auto_research_service.propose_topics()
    assert len(topics) >= 3, f"议题数不足: {len(topics)}"
    for t in topics:
        assert t["skills"], f"议题 {t['id']} skills 为空"
        for s in t["skills"]:
            assert s["id"] in SKILL_REGISTRY, f"非法算子 id: {s['id']}"
    print(f"共 {len(topics)} 个议题，全部通过校验。标题清单：")
    for t in topics:
        print(f"  [{t['id']}]（可行性:{t['feasibility']}）{t['title']}")
        print(f"      问题: {t['question']}")
        print(f"      算子: {[s['id'] for s in t['skills']]}")

    # ========== ② 完整流水线 ==========
    step("② start_pipeline 跑完整流水线（第一个议题）")
    topic = topics[0]
    job_id = auto_research_service.start_pipeline(topic)
    print(f"job_id = {job_id}，轮询中（超时 {TIMEOUT}s）...")

    t0 = time.time()
    while True:
        job = auto_research_service.get_job(job_id)
        if job["state"] in ("done", "failed"):
            break
        if time.time() - t0 > TIMEOUT:
            print("超时！当前 steps：")
            print(json.dumps(job["steps"], ensure_ascii=False, indent=1))
            sys.exit(1)
        running = [s["label"] for s in job["steps"] if s["state"] == "running"]
        print(f"  [{int(time.time() - t0):>4d}s] state={job['state']} running={running}", flush=True)
        time.sleep(10)

    print(f"\n流水线终态: {job['state']}, error={job.get('error')}")
    assert job["state"] == "done", f"流水线未完成: {job.get('error')}"

    # 断言：每步有终态
    for s in job["steps"]:
        assert s["state"] in ("done", "failed"), f"步骤 {s['key']} 无终态: {s['state']}"
    # 断言：result_ids ≥ 1
    assert len(job["result_ids"]) >= 1, "result_ids 为空"
    # 断言：paper / download_url 非空
    assert job["paper"] and job["paper"].get("title"), "paper 为空"
    assert job["download_url"], "download_url 为空"
    print(f"result_ids({len(job['result_ids'])}): {job['result_ids']}")
    print(f"论文标题: {job['paper']['title']}")
    print(f"download_url: {job['download_url']}")

    # 用 python-docx 读回验证六章节齐全
    fp = Path("data/outputs") / job["filename"]
    assert fp.exists(), f"论文文件不存在: {fp}"
    from docx import Document
    doc = Document(str(fp))
    text = "\n".join(p.text for p in doc.paragraphs)
    missing = [sec for sec in PAPER_SECTIONS if sec not in text]
    assert not missing, f"docx 缺章节: {missing}"
    print(f"docx 校验通过：{fp.name}，六章节齐全，段落数 {len(doc.paragraphs)}")

    # ========== ③ steps 时间线 ==========
    step("③ 流水线 steps 时间线")
    print(f"job {job['job_id']}  创建于 {job['created_at']}  完成于 {job['finished_at']}")
    for i, s in enumerate(job["steps"]):
        mark = {"done": "✓", "failed": "✗", "running": "…", "pending": "○"}[s["state"]]
        print(f"  {mark} [{i}] {s['label']:<14} {s['state']:<7} {s['detail']}"
              + (f"  (result_id={s['result_id']})" if s["result_id"] else ""))

    print("\n全部断言通过 ✓")


if __name__ == "__main__":
    main()
