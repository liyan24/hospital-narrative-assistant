"""
生成《前端功能与技术点演进文档》到 docs/ 目录
"""
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0, 102, 153)
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 102, 153)
    return heading


def add_paragraph(doc, text, bold=False, italic=False, indent=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    """添加项目符号"""
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)
    return p


def add_numbered(doc, text):
    """添加编号列表"""
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)
    return p


def main():
    doc = Document()

    # 设置默认中文字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)

    # 封面标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('医院叙事生成助手\n前端功能与技术点演进文档')
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'（从 Streamlit 迁移至 React 前端至今）\n\n生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # 1. 前言
    add_heading(doc, '一、前言', 1)
    add_paragraph(doc, '医院叙事生成助手是一个基于大语言模型（LLM）与 Neo4j 医疗知识图谱双引擎驱动的医院科室历史数据智能分析与叙事生成平台。项目早期使用 Streamlit 构建快速原型页面，随着功能复杂度和交互体验要求的提升，团队决定放弃 Streamlit 导航式页面，重新构建一套基于 React + Vite + Ant Design + ECharts 的现代化单页应用（SPA）前端。')
    add_paragraph(doc, '本文档系统梳理了从 Streamlit 前端迁移/重建开始，到目前为止（2026 年 6 月）所有主要功能迭代、技术点改造与问题修复，供团队复盘、交接与后续维护参考。')

    # 2. 前端架构演进
    add_heading(doc, '二、前端架构演进', 1)
    add_heading(doc, '2.1 技术栈选型', 2)
    add_paragraph(doc, '新前端采用以下技术栈：')
    add_bullet(doc, 'React 18：组件化 UI 框架')
    add_bullet(doc, 'Vite：下一代前端构建工具，提供极速的 HMR 与构建体验')
    add_bullet(doc, 'React Router v6：前端路由管理')
    add_bullet(doc, 'Ant Design 5：企业级 UI 组件库')
    add_bullet(doc, 'Ant Design Icons：图标库')
    add_bullet(doc, 'ECharts 5 + echarts-for-react：数据可视化')
    add_bullet(doc, 'echarts-wordcloud：词云图扩展')
    add_bullet(doc, 'Axios：HTTP 请求库')
    add_bullet(doc, 'Zustand（通过 stores/auth.jsx）：轻量级状态管理，用于认证与权限')
    add_bullet(doc, 'Day.js：日期处理')

    add_heading(doc, '2.2 项目结构', 2)
    add_paragraph(doc, 'frontend/ 目录结构如下：')
    add_bullet(doc, 'src/views/portal/：医生前台视图（工作台、患者全息视图、查房助手、晨会简报、相似患者/诊疗路径、质控闭环管理、科研队列等）')
    add_bullet(doc, 'src/views/admin/：后台管理视图（用户管理、角色权限、功能开关、系统配置、缓存管理、概览仪表盘）')
    add_bullet(doc, 'src/layouts/：PortalLayout（医生前台布局）与 AdminLayout（后台管理布局）')
    add_bullet(doc, 'src/api/index.js：统一封装的 Axios 请求')
    add_bullet(doc, 'src/components/：可复用组件，如 PatientSearch、PatientIdLink')
    add_bullet(doc, 'src/stores/auth.jsx：认证状态与权限/功能开关管理')
    add_bullet(doc, 'src/router/index.jsx：前端路由配置')

    # 3. 功能模块演进
    add_heading(doc, '三、功能模块演进', 1)

    add_heading(doc, '3.1 医生前台（Portal）', 2)

    add_heading(doc, '3.1.1 工作台', 3)
    add_bullet(doc, '登录后默认首页，展示系统入口与简要说明')
    add_bullet(doc, '顶部导航菜单根据用户角色与功能开关动态渲染')

    add_heading(doc, '3.1.2 患者全息视图（PatientDetailView）', 3)
    add_paragraph(doc, '患者全息视图是前台的核心模块之一，经历了多轮优化：')
    add_numbered(doc, '图文时间轴精简：从最初的完整就诊列表，精简为只展示关键事件（手术、检查），避免数据量大时页面冗长。')
    add_numbered(doc, '就诊概览卡片：新增按次就诊的卡片式概览，展示日期、住院天数、主诉、诊断、手术/用药/检查数量。')
    add_numbered(doc, '住院趋势图：使用 ECharts 折线图展示历次住院天数。')
    add_numbered(doc, '住院情况概览（最新）：将关键事件时间轴与住院趋势图合二为一，折线图 X 轴以第一次住院为起点，并在折线上用 markPoint 标记每次住院的检查/检验数量，tooltip 固定宽度、鼠标右侧显示。')
    add_numbered(doc, '就诊概览词云图（最新）：将卡片列表替换为“诊断词云”和“用药词云”两张 ECharts 词云图，直观展示高频诊断与用药。')
    add_numbered(doc, '关键异常提示：基于风险评分、质控问题、住院天数、化疗/手术史、再入院次数等生成异常标签。')
    add_numbered(doc, '智能病程摘要：展示基于知识图谱生成的患者完整叙事。')
    add_numbered(doc, '风险预测与质控：展示风险等级、风险评分、风险因素及质控问题列表。')
    add_numbered(doc, '再入院分析： originally 与相似患者合并展示，后拆分为独立 Card，并去除 200 字截断，默认完整展示，超长时滚动。')
    add_numbered(doc, '相似患者：展示 Top 5 相似患者及相似度。')
    add_numbered(doc, '患者 ID 链接组件化：所有患者 ID 均使用 PatientIdLink 组件，点击可跳转至全息视图。')

    add_heading(doc, '3.1.3 查房助手（WardRoundView）', 3)
    add_bullet(doc, '支持选择患者后基于知识图谱 RAG 回答医疗问题')
    add_bullet(doc, '修复了 RAG 问答中患者 ID 识别问题：KGRAGRequest 增加 patient_id 字段，后端优先使用传入 patient_id，而非依赖文本解析')
    add_bullet(doc, '可基于真实就诊记录回答，例如 51 次就诊的患者')

    add_heading(doc, '3.1.4 科室晨会简报（BriefingView）', 3)
    add_bullet(doc, '修复初始化日期问题：直接使用模拟日期加载，避免 setDate 异步导致首次渲染为空')
    add_bullet(doc, '支持按日期查询每日简报')
    add_bullet(doc, '根据角色权限展示科室选择器')

    add_heading(doc, '3.1.5 相似患者 / 诊疗路径（SimilarPatientView）', 3)
    add_bullet(doc, '相似患者推荐：输入患者 ID 后查找相似病例，展示相似度、共同诊断、共同用药')
    add_bullet(doc, '修复相似度显示 NaN% 的问题：前端表格 dataIndex 由 similarity 改为后端返回的 score')
    add_bullet(doc, '优化患者搜索框宽度：PatientSearch 组件默认占满 100% 宽度，与诊疗路径推荐的 Input 样式保持一致')
    add_bullet(doc, '诊疗路径推荐：输入疾病名称后生成诊疗路径叙事与关键数据')

    add_heading(doc, '3.1.6 质控闭环管理（QualityView）', 3)
    add_bullet(doc, '三个 Tab：今日质控异常、患者质控问题、整改追踪')
    add_bullet(doc, '修复字段名不匹配：将 briefing.quality_issues 改为 briefing.quality_control_issues')
    add_bullet(doc, '修复患者质控问题不显示：后端 issues 为数组，前端原按对象遍历导致为空，改为直接遍历数组')
    add_bullet(doc, '选择患者后自动触发查询')
    add_bullet(doc, '操作列仅对 admin / department_manager / quality_controller 开放，可修改整改状态')

    add_heading(doc, '3.1.7 科研队列', 3)
    add_bullet(doc, '支持按疾病、用药、检查等条件筛选患者队列')
    add_bullet(doc, '患者 ID 使用 PatientIdLink 组件')

    add_heading(doc, '3.2 后台管理（Admin）', 2)
    add_bullet(doc, '用户管理：新增/编辑/删除用户，支持角色分配')
    add_bullet(doc, '角色权限管理：新增角色权限编辑页面，支持 9 角色体系')
    add_bullet(doc, '功能开关：动态控制前台菜单与功能是否展示')
    add_bullet(doc, '系统配置：包含模拟日期等配置')
    add_bullet(doc, '缓存管理：展示 LLM 缓存统计，支持清理缓存')
    add_bullet(doc, '后台概览仪表盘：展示知识图谱统计、LLM 缓存统计、系统状态')
    add_bullet(doc, '修复后台概览页 KG 节点/关系数为 0 的问题：前端适配后端返回的 nodes/relationships 结构')
    add_bullet(doc, '修复 LLM 缓存无条目显示问题：前端读取 cacheStats.stats 字段')

    # 4. 技术点与改造
    add_heading(doc, '四、关键技术点与后端改造', 1)

    add_heading(doc, '4.1 角色权限与功能开关体系', 2)
    add_paragraph(doc, '建立了 9 角色体系：admin、department_manager、hospital_manager、quality_controller、attending_doctor、resident_doctor、director、researcher、viewer。菜单根据 permissions 数组与 features 功能开关动态渲染。')

    add_heading(doc, '4.2 认证体系', 2)
    add_bullet(doc, '后端使用 bcrypt 加密密码、PyJWT 签发 Token')
    add_bullet(doc, '前端使用 localStorage 存储 token，Axios 请求头携带 Authorization')
    add_bullet(doc, '后台路由单独鉴权')

    add_heading(doc, '4.3 LLM 缓存机制', 2)
    add_bullet(doc, 'database/llm_cache.py 实现基于文件系统的 LLM 输出缓存')
    add_bullet(doc, '缓存键基于 namespace + 内容哈希，避免重复调用 LLM')
    add_bullet(doc, '支持 TTL（默认 240 小时）、过期自动清理、按命名空间清理')
    add_bullet(doc, 'llm_service.chat 自动读取/写入缓存')

    add_heading(doc, '4.4 结果缓存机制（新增）', 2)
    add_bullet(doc, '新增 database/result_cache.py 通用结果缓存模块')
    add_bullet(doc, '为风险预测、相似患者、再入院患者叙事等服务添加完整结果缓存')
    add_bullet(doc, '显著降低患者全息视图加载时间：相似患者从 10 秒级降至 10 毫秒级')

    add_heading(doc, '4.5 MySQLClient 修复', 2)
    add_bullet(doc, '修复 INSERT/UPDATE/DELETE 不支持 lastrowid 的问题')
    add_bullet(doc, '修复用户 CRUD 和角色权限更新 500 错误')

    add_heading(doc, '4.6 依赖补充', 2)
    add_bullet(doc, 'pyproject.toml 补充 bcrypt、PyJWT、openpyxl')
    add_bullet(doc, '前端补充 echarts-wordcloud')

    add_heading(doc, '4.7 模拟日期配置', 2)
    add_bullet(doc, 'config.py 新增 simulation_date 配置，用于测试环境固定日期')
    add_bullet(doc, '每日简报默认使用模拟日期，避免首次渲染为空')

    # 5. 问题修复记录
    add_heading(doc, '五、重点问题修复记录', 1)

    add_heading(doc, '5.1 前端显示类', 2)
    add_bullet(doc, '后台概览页 KG 统计为 0：前端字段名不匹配，已适配')
    add_bullet(doc, '后台概览页 LLM 缓存无条目：前端未取 stats 嵌套字段，已修复')
    add_bullet(doc, '相似患者相似度 NaN%：前端使用 similarity 字段而后端返回 score，已统一')
    add_bullet(doc, '再入院分析 500：readmission_service.py 漏引入 result_cache_store，已补充 import')
    add_bullet(doc, '质控闭环管理无数据：字段名 quality_issues 与 quality_control_issues 不匹配，已修复')
    add_bullet(doc, '患者质控问题不显示：前端将数组误按对象遍历，已修正')
    add_bullet(doc, '再入院叙事截断：原代码手动 slice(0,200)，已改为完整展示')

    add_heading(doc, '5.2 性能优化类', 2)
    add_bullet(doc, '患者全息视图三个板块（风险预测、再入院、相似患者）添加结果缓存，二次加载从秒级降至毫秒级')
    add_bullet(doc, '质控、再入院、风险预测等服务均接入 LLM 缓存')

    add_heading(doc, '5.3 交互优化类', 2)
    add_bullet(doc, '患者全息视图时间轴与住院趋势合并为“住院情况概览”')
    add_bullet(doc, '就诊概览改为诊断/用药词云')
    add_bullet(doc, '再入院分析与相似患者拆分为两个独立区域')
    add_bullet(doc, '智能病程摘要与再入院分析固定等高 520px，默认展开')
    add_bullet(doc, 'PatientSearch 搜索框默认占满宽度')

    # 6. 当前状态
    add_heading(doc, '六、当前运行状态', 1)
    add_bullet(doc, '后端 FastAPI 运行于 http://127.0.0.1:8005')
    add_bullet(doc, '前端 Vite dev server 运行于 http://localhost:8501')
    add_bullet(doc, '知识图谱：约 35,000+ 节点、780,000+ 关系')
    add_bullet(doc, 'LLM 缓存：20+ 命名空间，缓存大量报告、RAG、风险、相似患者等结果')
    add_bullet(doc, '数据模拟日期：2025-05-29')

    # 7. 后续建议
    add_heading(doc, '七、后续优化建议', 1)
    add_numbered(doc, '为后台任务超时问题提供进程守护或持久化部署方案（当前平台限制 1 小时超时）。')
    add_numbered(doc, '为结果缓存增加手动刷新/失效机制，避免数据更新后仍读取旧缓存。')
    add_numbered(doc, '为 ECharts tooltip、词云图等可视化细节增加更多自适应与移动端适配。')
    add_numbered(doc, '补充前端单元测试与端到端测试，覆盖核心交互流程。')
    add_numbered(doc, '生产环境修改默认密码与 JWT Secret。')

    # 保存文档
    docs_dir = Path(__file__).parent.parent / 'docs'
    docs_dir.mkdir(exist_ok=True)
    filename = docs_dir / f'前端功能与技术点演进文档_{datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(filename)
    print(f'文档已生成：{filename}')


if __name__ == '__main__':
    main()
