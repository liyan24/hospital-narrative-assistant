"""
周简报文档生成服务：生成Word和PDF格式的周简报。
"""
import os
import io
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt


class WeeklyDocumentService:
    def __init__(self):
        self.output_dir = Path("./data/outputs")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._register_fonts()

    def _register_fonts(self):
        font_paths = [
            ("C:/Windows/Fonts/simhei.ttf", "SimHei"),
            ("C:/Windows/Fonts/simsun.ttc", "SimSun"),
            ("C:/Windows/Fonts/msyh.ttc", "MicrosoftYaHei"),
        ]
        for fp, name in font_paths:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont(name, fp))
                except Exception:
                    pass

    def _get_cn_font(self):
        registered = pdfmetrics.getRegisteredFontNames()
        for name in ["SimHei", "SimSun", "MicrosoftYaHei", "Helvetica"]:
            if name in registered:
                return name
        return "Helvetica"

    def _add_markdown_paragraph(self, doc, text: str):
        """在docx中添加段落，解析Markdown格式（**加粗**、*斜体*）"""
        if not text:
            return doc.add_paragraph()
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                inner = part[2:-2]
                inner_parts = re.split(r'(\*.*?\*)', inner)
                for ip in inner_parts:
                    if ip.startswith('*') and ip.endswith('*') and len(ip) > 1:
                        run = p.add_run(ip[1:-1])
                        run.bold = True
                        run.italic = True
                    else:
                        run = p.add_run(ip)
                        run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 1:
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                p.add_run(part)
        return p

    def _add_table_to_docx(self, doc, headers, rows):
        """在docx中添加表格"""
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # 表头
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = str(h)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        # 数据行
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                table.rows[r_idx + 1].cells[c_idx].text = str(val)
                for paragraph in table.rows[r_idx + 1].cells[c_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
        return table

    def generate_weekly_docx(self, report: dict) -> str:
        """生成周简报Word文档"""
        filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_周简报.docx"
        filepath = self.output_dir / filename

        doc = Document()
        data = report.get("data", {})
        texts = report.get("texts", {})
        week_range = data.get("week_range", "")

        # 标题
        title = doc.add_heading("肿瘤血液科 · 每周临床简报", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"数据周期：{week_range}", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"生成时间：{data.get('generated_at', '')}")
        doc.add_paragraph()

        # 模块1：本周运营概况
        doc.add_heading("模块1：本周运营概况", level=1)
        op = data.get("operation", {})
        # 核心运营指标表
        self._add_table_to_docx(doc,
            ["指标", "本周数据", "上周环比", "去年同期同比"],
            [
                ["入院人次", f"{op.get('admission_count', 0)}人", self._fmt_change(op.get('admission_vs_prev')), self._fmt_change(op.get('admission_vs_yoy'))],
                ["出院人次", f"{op.get('discharge_count', 0)}人", self._fmt_change(op.get('discharge_vs_prev')), self._fmt_change(op.get('discharge_vs_yoy'))],
                ["在院患者日均", f"{op.get('avg_in_hospital', 0)}人", "-", "-"],
                ["床位使用率", f"{op.get('bed_usage_rate', 0)}%", "-", "-"],
                ["平均住院天数", f"{op.get('avg_hospitalization_days', 0)}天", "-", "-"],
            ]
        )
        doc.add_paragraph()
        self._add_markdown_paragraph(doc, texts.get("operation", ""))
        doc.add_paragraph()

        # 每日入院分布表
        daily = op.get("daily_admission", {})
        self._add_table_to_docx(doc,
            ["星期", "周一", "周二", "周三", "周四", "周五", "周六", "周日", "合计"],
            [
                ["入院人数"] + [daily.get(f"周{n}", 0) for n in ["一", "二", "三", "四", "五", "六", "日"]] + [op.get("admission_count", 0)],
            ]
        )
        doc.add_paragraph()

        # 模块2：病种分析
        doc.add_heading("模块2：病种分析", level=1)
        diseases = data.get("diseases", {})
        top5 = diseases.get("top5", [])
        if top5:
            self._add_table_to_docx(doc,
                ["排名", "病种", "本周入院", "占比", "环比", "累计年度占比"],
                [[i+1, d["disease"], f"{d['count']}人", f"{d['percentage']}%", "-", "-"] for i, d in enumerate(top5)]
            )
        doc.add_paragraph()
        self._add_markdown_paragraph(doc, texts.get("diseases", ""))
        doc.add_paragraph()

        # 模块3：检查检验汇总
        doc.add_heading("模块3：检查检验汇总", level=1)
        exam_lab = data.get("exam_lab", {})
        exam_types = exam_lab.get("exam_types", [])
        if exam_types:
            self._add_table_to_docx(doc,
                ["检查类型", "数量", "占比", "阳性/异常", "阳性率"],
                [[e["type"], f"{e['count']}例", "-", f"{e['positive']}例", f"{e['positive_rate']}%"] for e in exam_types[:5]]
            )
        doc.add_paragraph()
        self._add_markdown_paragraph(doc, texts.get("exam_lab", ""))
        doc.add_paragraph()

        # 模块4：治疗动态
        doc.add_heading("模块4：治疗动态", level=1)
        treatment = data.get("treatment", {})
        surgeries = treatment.get("surgeries", [])
        if surgeries:
            self._add_table_to_docx(doc,
                ["日期", "手术名称", "类型"],
                [[s["date"], s["name"], s["type"]] for s in surgeries]
            )
        doc.add_paragraph()
        self._add_markdown_paragraph(doc, texts.get("treatment", ""))
        doc.add_paragraph()

        # 模块5：质控指标
        doc.add_heading("模块5：质控指标", level=1)
        quality = data.get("quality", {})
        self._add_table_to_docx(doc,
            ["质控指标", "本周", "目标值", "状态"],
            [
                ["平均住院天数", f"{quality.get('avg_days', 0)}天", "≤10天", "达标" if quality.get('avg_days', 0) <= 10 else "未达标"],
                ["30天非计划再入院率", f"{quality.get('readmit_30_rate', 0)}%", "≤5%", "达标" if quality.get('readmit_30_rate', 0) <= 5 else "未达标"],
            ]
        )
        doc.add_paragraph()
        self._add_markdown_paragraph(doc, texts.get("quality", ""))
        doc.add_paragraph()

        # 模块6：重点关注患者
        doc.add_heading("模块6：重点关注患者", level=1)
        focus = data.get("focus_patients", {})
        elderly = focus.get("elderly", [])
        if elderly:
            doc.add_paragraph(f"一、本周新入院高龄患者（≥80岁，共{focus.get('elderly_count', 0)}人）")
            self._add_table_to_docx(doc,
                ["姓名", "性别/年龄", "诊断"],
                [[e["name"], e["gender_age"], e["diagnosis"]] for e in elderly]
            )
        long_stay = focus.get("long_stay", [])
        if long_stay:
            doc.add_paragraph(f"二、本周超长住院患者（>30天，共{focus.get('long_stay_count', 0)}人）")
            self._add_table_to_docx(doc,
                ["姓名", "诊断", "住院天数"],
                [[l["name"], l["diagnosis"], f"{l['days']}天"] for l in long_stay]
            )
        doc.add_paragraph()
        self._add_markdown_paragraph(doc, texts.get("focus_patients", ""))
        doc.add_paragraph()

        # 模块7：下周预警
        doc.add_heading("模块7：下周预警", level=1)
        next_week = data.get("next_week", {})
        doc.add_paragraph(f"预测周期：{next_week.get('next_week_range', '')}")
        doc.add_paragraph(f"预计入院总量：{next_week.get('forecast_admission', 0)}人次")
        doc.add_paragraph(f"预计手术：{next_week.get('forecast_surgeries', 0)}台")
        doc.add_paragraph()
        self._add_markdown_paragraph(doc, texts.get("next_week", ""))
        doc.add_paragraph()

        # 附录
        doc.add_heading("附录", level=1)
        doc.add_paragraph("报告说明：本简报数据来源于医院信息系统（HIS）、实验室信息系统（LIS）、影像归档系统（PACS）及电子病历系统（EMR），数据统计截止至每周日24:00。")

        doc.save(filepath)
        return str(filepath)

    def _fmt_change(self, val):
        if val is None:
            return "-"
        if val > 0:
            return f"↑+{val}%"
        elif val < 0:
            return f"↓{val}%"
        return "持平"

    def _markdown_to_html(self, text: str) -> str:
        """将Markdown格式转换为简单的HTML标签（用于ReportLab）"""
        if not text:
            return text
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
        return text

    def generate_weekly_pdf(self, report: dict) -> str:
        """生成周简报PDF文档（简化版）"""
        filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_周简报.pdf"
        filepath = self.output_dir / filename

        font_name = self._get_cn_font()
        styles = getSampleStyleSheet()
        cn_style = ParagraphStyle("Chinese", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=14, spaceAfter=4)
        title_style = ParagraphStyle("ChineseTitle", parent=styles["Title"], fontName=font_name, fontSize=18, alignment=1, spaceAfter=12)
        h1_style = ParagraphStyle("ChineseH1", parent=styles["Heading1"], fontName=font_name, fontSize=14, spaceAfter=8)

        doc = SimpleDocTemplate(str(filepath), pagesize=A4,
            rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
        story = []

        data = report.get("data", {})
        texts = report.get("texts", {})
        week_range = data.get("week_range", "")

        story.append(Paragraph("肿瘤血液科 · 每周临床简报", title_style))
        story.append(Paragraph(f"数据周期：{week_range}", cn_style))
        story.append(Spacer(1, 0.5*cm))

        sections = [
            ("模块1：本周运营概况", "operation"),
            ("模块2：病种分析", "diseases"),
            ("模块3：检查检验汇总", "exam_lab"),
            ("模块4：治疗动态", "treatment"),
            ("模块5：质控指标", "quality"),
            ("模块6：重点关注患者", "focus_patients"),
            ("模块7：下周预警", "next_week"),
        ]

        for section_title, section_key in sections:
            story.append(Paragraph(section_title, h1_style))
            text = texts.get(section_key, "")
            for line in text.split("\n"):
                stripped = line.strip()
                if stripped:
                    story.append(Paragraph(self._markdown_to_html(stripped), cn_style))
                else:
                    story.append(Spacer(1, 0.2*cm))
            story.append(Spacer(1, 0.3*cm))

        doc.build(story)
        return str(filepath)

    def export_weekly(self, report: dict, fmt: str = "docx") -> str:
        if fmt.lower() == "docx":
            return self.generate_weekly_docx(report)
        elif fmt.lower() == "pdf":
            return self.generate_weekly_pdf(report)
        else:
            raise ValueError(f"不支持的格式: {fmt}")


# 全局单例
weekly_document_service = WeeklyDocumentService()
