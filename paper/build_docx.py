#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""把 paper/paper_zh.md 与 paper/paper_en.md 组装为两个 Word 文档（含 4 张插图）。

可复跑：直接覆盖输出文件。不修改任何源文件。
用法：.venv/Scripts/python.exe paper/build_docx.py
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor

if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parent
FIG_DIR = BASE / "figures"

# (md 文件, 输出 docx, 图题语言)
JOBS = [
    (BASE / "paper_zh.md", BASE / "论文_中文版_v1.docx", "zh"),
    (BASE / "paper_en.md", BASE / "paper_english_v1.docx", "en"),
]

# 插图配置：所在小节号 -> (图片文件名, 中文图题, 英文图题, 正文触发词)
FIGURES = {
    "3.3":   ("figure1_framework.png",
              "图 1  知识图谱增强叙事生成总体框架",
              "Figure 1. Overall framework of KG-grounded narrative generation",
              ("图 1", "Figure 1")),
    "3.2.3": ("figure2_ontology.png",
              "图 2  知识图谱本体 schema（9 类节点 + 9 类关系）",
              "Figure 2. Knowledge graph ontology schema (9 node types, 9 relation types)",
              ("图 2", "Figure 2")),
    "3.4.3": ("figure3_evaluation.png",
              "图 3  三重评估协议流程（自动核查 / LLM 评审 / 盲法评审）",
              "Figure 3. Three-layer evaluation protocol (automated verification / LLM-as-Judge / blinded review)",
              ("图 3", "Figure 3")),
    "4.6":   ("figure4_case_study.png",
              "图 4  案例研究对比图（B2 Vector-RAG 对 B3，标注幻觉与溯源链）",
              "Figure 4. Case study comparison (B2 Vector-RAG vs B3)",
              ("图 4", "Figure 4")),
}

BODY_SIZE = Pt(12)        # 小四
TABLE_SIZE = Pt(10.5)     # 五号
CAPTION_SIZE = Pt(9)      # 小五


# ---------- 字体工具 ----------

def set_run_font(run, size=BODY_SIZE, bold=None, italic=None,
                 ascii_font="Times New Roman", ea_font="宋体"):
    run.font.name = ascii_font
    run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), ea_font)


def shade_paragraph(paragraph, fill="F2F2F2"):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


# ---------- 行内格式解析 ----------

def strip_latex(text):
    """去掉 $ 与常见 LaTeX 记号，按纯文本输出。"""
    text = re.sub(r"\$([^$]*)\$", lambda m: m.group(1), text)  # 行内 $...$
    text = text.replace("$$", "")
    text = re.sub(r"\\text\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\sqrt\{([^}]*)\}", r"√(\1)", text)
    text = text.replace("\\{", "{").replace("\\}", "}")
    text = text.replace("\\alpha", "α").replace("\\kappa", "κ")
    text = text.replace("\\", "")  # 其余反斜杠直接去掉
    return text


# 行内记号：代码 > 加粗 > 斜体。
# 定界符前面不允许是 ASCII 字母数字或 *（避免把 "p < 0.001**" 这类残留的 **
# 误当定界符；注意不能用 \w，否则中文相邻的 *斜体* 无法匹配）。
_TOKEN_RE = re.compile(
    r"`[^`]+`"
    r"|(?<![A-Za-z0-9*])\*\*(?=\S).+?(?<=\S)\*\*"
    r"|(?<![A-Za-z0-9*])\*(?!\*)(?=\S).+?(?<=\S)\*(?![A-Za-z0-9*])"
)


def add_inline_runs(paragraph, text, size=BODY_SIZE, base_italic=False):
    """把含 Markdown 行内记号的文本写入段落，应用加粗/斜体/等宽字体。"""
    pos = 0
    for m in _TOKEN_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            set_run_font(r, size=size, italic=base_italic or None)
        tok = m.group(0)
        if tok.startswith("`"):
            r = paragraph.add_run(tok[1:-1])
            set_run_font(r, size=size, italic=base_italic or None,
                         ascii_font="Consolas", ea_font="宋体")
        elif tok.startswith("**"):
            r = paragraph.add_run(tok[2:-2])
            set_run_font(r, size=size, bold=True, italic=base_italic or None)
        else:
            r = paragraph.add_run(tok[1:-1])
            set_run_font(r, size=size, italic=True)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_run_font(r, size=size, italic=base_italic or None)


# ---------- 文档构建 ----------

HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")
SECTION_NO_RE = re.compile(r"^(\d+(?:\.\d+)*)")


def new_document():
    doc = Document()
    # A4 页面
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    # 正文默认字体：英文 Times New Roman / 中文宋体，小四
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return doc


def add_heading(doc, text, level):
    h = doc.add_heading("", level=level)
    r = h.add_run(text)
    set_run_font(r, size=Pt({1: 16, 2: 14, 3: 13, 4: 12}[level]),
                 bold=True, ascii_font="Arial", ea_font="黑体")
    r.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_table(doc, rows):
    """rows: list of list of str；第一行为表头。"""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, cell_text, size=TABLE_SIZE)
            if i == 0:
                for r in p.runs:
                    r.font.bold = True
    return table


def parse_table_block(lines):
    rows = []
    for ln in lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
            continue  # 分隔行
        rows.append(cells)
    return rows


def add_figure(doc, png_name, caption):
    path = FIG_DIR / png_name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(15.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=CAPTION_SIZE, bold=True)


def build(md_path, out_path, lang):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = new_document()

    current_section = None      # 当前小节号，如 "3.2.3"
    inserted = set()            # 已插入的小节号
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 表格块
        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            rows = parse_table_block(block)
            if rows:
                add_table(doc, rows)
            continue

        i += 1
        if not stripped or stripped == "---":
            continue

        # 标题
        m = HEADING_RE.match(stripped)
        if m:
            level = len(m.group(1))
            htext = strip_latex(m.group(2).strip())
            add_heading(doc, htext, level)
            sm = SECTION_NO_RE.match(htext)
            if sm:
                current_section = sm.group(1).rstrip(".")
            elif level <= 2:
                current_section = None
            continue

        # 引用块（草稿提示等，原样保留，浅灰底纹 + 斜体）
        if stripped.startswith(">"):
            content = strip_latex(stripped.lstrip(">").strip())
            p = doc.add_paragraph()
            shade_paragraph(p)
            add_inline_runs(p, content, base_italic=True)
            continue

        # 项目符号列表
        if stripped.startswith("- "):
            content = strip_latex(stripped[2:].strip())
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, content)
            continue

        # 普通段落
        content = strip_latex(stripped)
        p = doc.add_paragraph()
        add_inline_runs(p, content)

        # 按正文引用处插入插图
        if current_section in FIGURES and current_section not in inserted:
            png, cap_zh, cap_en, triggers = FIGURES[current_section]
            if any(t in content for t in triggers):
                add_figure(doc, png, cap_zh if lang == "zh" else cap_en)
                inserted.add(current_section)

    missing = set(FIGURES) - inserted
    if missing:
        print(f"  警告：以下插图未找到插入位置：{sorted(missing)}", file=sys.stderr)

    doc.save(out_path)
    return out_path


# ---------- 读回验证 ----------

def verify(docx_path, md_path):
    from docx import Document as D
    doc = D(docx_path)

    paragraphs = doc.paragraphs
    n_tables = len(doc.tables)
    n_images = len(doc.inline_shapes)
    headings = [(p.style.name, p.text) for p in paragraphs
                if p.style.name.startswith("Heading")]
    md_tables = sum(
        1 for ln in md_path.read_text(encoding="utf-8").splitlines()
        if ln.strip().startswith("|") and re.fullmatch(r"[\s|:>-]*", ln.strip())
    )

    print(f"\n=== {docx_path.name} ===")
    print(f"段落数: {len(paragraphs)}")
    print(f"表格数: {n_tables}（md 中表格数: {md_tables}）-> {'OK' if n_tables == md_tables else '不匹配!'}")
    print(f"图片数: {n_images} -> {'OK' if n_images == 4 else '不为 4!'}")
    from collections import Counter
    c = Counter(s for s, _ in headings)
    print(f"标题层级: {dict(sorted(c.items()))}")
    print("前几段文本抽查:")
    shown = 0
    for p in paragraphs:
        t = p.text.strip()
        if t:
            print(f"  [{p.style.name}] {t[:60]}")
            shown += 1
            if shown >= 8:
                break


def main():
    for md, out, lang in JOBS:
        path = build(md, out, lang)
        print(f"已生成: {path}")
    for md, out, _ in JOBS:
        verify(out, md)


if __name__ == "__main__":
    main()
