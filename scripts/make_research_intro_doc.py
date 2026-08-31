# -*- coding: utf-8 -*-
"""生成《科研助手功能简介（汇报版）》Word 文档（一次性工具）。"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
IMG = ROOT / "docs" / "images" / "research"
OUT = ROOT / "docs" / "科研助手功能简介（汇报版）.docx"


def set_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text, size=12, bold=False, font="宋体", indent=True,
             align=None, space_after=6):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    if align:
        p.alignment = align
    set_font(p.add_run(text), name=font, size=size, bold=bold)
    return p


def add_heading(doc, text, size=15):
    add_para(doc, text, size=size, bold=True, font="黑体", indent=False,
             space_after=8)


def add_image(doc, path, width_cm=15.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    if caption:
        add_para(doc, caption, size=9, indent=False,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)


doc = Document()

# 标题
add_para(doc, "科研助手功能简介", size=20, bold=True, font="黑体",
         indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, "——让医生不写代码也能做数据科研", size=13, font="楷体",
         indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

# 一、为什么做
add_heading(doc, "一、为什么做这个功能")
add_para(doc, "我们科室积累了大量真实诊疗数据（3990 名患者、1.37 万次住院记录、"
              "约 114 万条医嘱、5.5 万条检验结果），但医生普遍不具备数据挖掘和"
              "机器学习的技术手段，数据长期“沉睡”，难以转化为科研成果。")
add_para(doc, "科研助手的目标，就是把数据分析的技术门槛降到零：医生不需要学编程、"
              "不需要懂算法，点几下鼠标，就能完成从选题、数据分析到论文初稿的"
              "完整过程。")

# 二、它是什么
add_heading(doc, "二、它是什么")
add_para(doc, "科研助手是“医院叙事生成助手”系统中的一个模块。它把科研中常用的"
              "数据分析方法（关联规律挖掘、相似患者查找、分类预测、聚类分群、"
              "统计检验等 24 种）预先做成了现成的分析工具，再由 AI 大模型把"
              "这些工具串起来，自动完成整个科研流程。医生只需要做选择题和判断题，"
              "其余工作全部由系统自动完成。")

# 三、怎么用
add_heading(doc, "三、怎么用：三步出一篇论文")
add_para(doc, "第一步：推荐选题。系统自动评估科室数据情况后，推荐若干个"
              "“数据撑得住”的研究题目，并说明推荐理由和打算用什么方法分析。"
              "如果医生有自己的想法，也可以直接输入，系统会评估现有数据是否"
              "支持、并给出修改建议。", indent=False)
add_para(doc, "第二步：一键研究。选定题目后点击开始，系统自动执行全套分析"
              "（合并症规律、影响因素、风险预测等），每一步进度实时可见，"
              "全程约 3～6 分钟，期间不需要任何操作。", indent=False)
add_para(doc, "第三步：下载论文。系统自动生成符合医学论文规范的中文初稿"
              "（摘要、前言、方法、结果、讨论、结论、参考文献），"
              "以 Word 文档下载，可直接修改。所有做过的研究都保存在"
              "历史记录中，随时可以查看原文或重复下载。", indent=False)

add_image(doc, IMG / "01_智能模式_议题推荐.png",
          caption="图 1　系统根据科室真实数据自动推荐研究选题")
add_image(doc, IMG / "03_智能模式_流水线运行.png",
          caption="图 2　自动研究过程实时展示，医生无需任何操作")

# 四、能做什么分析
add_heading(doc, "四、它能做哪些分析")
add_para(doc, "找规律：哪些疾病经常同时出现（合并症模式），哪些药物经常联合使用；", indent=False)
add_para(doc, "找因素：哪些临床特征与患者再入院、住院时间延长显著相关；", indent=False)
add_para(doc, "做预测：预测患者再入院风险、住院天数，并指出最重要的影响因素；", indent=False)
add_para(doc, "分人群：把患者按病情特征自动分组，比较不同组的特点；", indent=False)
add_para(doc, "查文献、写论文：自动检索 PubMed 文献、拟定提纲、撰写初稿、"
              "整理参考文献格式。", indent=False)

# 五、实际效果
add_heading(doc, "五、实际效果举例")
add_para(doc, "以真实运行的一个题目“老年肺部感染患者抗菌药物联合用药模式分析”"
              "为例：系统在几分钟内自动完成了对 1.3 万例住院记录的分析，"
              "发现了“肺大疱→肺气肿”等合并症关联（置信度 79.8%），"
              "建立的再入院预测模型准确率达 83.5%，并自动撰写了结构完整的"
              "论文初稿（图 3）。分析过程可复现、结果可核查。")
add_image(doc, IMG / "07_报告_docx_第1页.png", width_cm=13.5,
          caption="图 3　系统自动生成的论文初稿首页")

# 六、使用注意
add_heading(doc, "六、需要说明的几点")
add_para(doc, "1. 论文是 AI 起草的初稿，其中所有数据结论都来自系统对真实数据的"
              "实际计算，但投稿前必须由作者人工审校，并补齐伦理审批等环节；", indent=False)
add_para(doc, "2. 文献检索功能需要连接外网；", indent=False)
add_para(doc, "3. 部分数据覆盖有限（如检验数据仅约两成患者有），系统推荐选题时"
              "会自动避开这些短板，保证研究可行。", indent=False)

doc.save(str(OUT))
print("saved:", OUT)
