"""
文档生成服务：支持导出Word和PDF报告（含图表）。
"""
import os
import io
import base64
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import re


class DocumentService:
    def __init__(self):
        self.output_dir = Path("./data/outputs")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._register_fonts()

    def _register_fonts(self):
        """注册中文字体（ReportLab需要）"""
        font_paths = [
            ("C:/Windows/Fonts/simhei.ttf", "SimHei"),
            ("C:/Windows/Fonts/simsun.ttc", "SimSun"),
            ("C:/Windows/Fonts/msyh.ttc", "MicrosoftYaHei"),
            ("C:/Windows/Fonts/msyhbd.ttc", "MicrosoftYaHeiBold"),
        ]
        for fp, name in font_paths:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont(name, fp))
                except Exception:
                    pass

    def _get_cn_font(self):
        registered = pdfmetrics.getRegisteredFontNames()
        for name in ["SimHei", "SimSun", "MicrosoftYaHei", "Helvetica"]:
            if name in registered:
                return name
        return "Helvetica"

    def _markdown_to_html(self, text: str) -> str:
        """将Markdown格式转换为简单的HTML标签（用于ReportLab）"""
        if not text:
            return text
        # 加粗 **text**
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        # 斜体 *text*（但排除已处理的<b>内部）
        text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
        return text

    def _add_markdown_paragraph(self, doc, text: str):
        """在docx中添加段落，解析Markdown格式（**加粗**、*斜体*）"""
        if not text:
            return doc.add_paragraph()
        p = doc.add_paragraph()
        # 按 ** 分割，处理加粗
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # 加粗文本
                inner = part[2:-2]
                # 内部可能还有斜体
                inner_parts = re.split(r'(\*.*?\*)', inner)
                for ip in inner_parts:
                    if ip.startswith('*') and ip.endswith('*') and len(ip) > 1:
                        run = p.add_run(ip[1:-1])
                        run.bold = True
                        run.italic = True
                    else:
                        run = p.add_run(ip)
                        run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 1:
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                p.add_run(part)
        return p

    def _chart_to_image(self, chart_config: dict, width: int = 8, height: int = 5) -> bytes:
        """将简化图表配置用matplotlib渲染为PNG字节"""
        import matplotlib
        matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial Unicode MS', 'DejaVu Sans']
        matplotlib.rcParams['axes.unicode_minus'] = False

        fig, ax = plt.subplots(figsize=(width, height))
        title = chart_config.get("title", {}).get("text", "")
        series = chart_config.get("series", [])
        x_data = chart_config.get("xAxis", {}).get("data", [])

        if not series:
            plt.close()
            return b""

        for s in series:
            s_type = s.get("type", "bar")
            data = s.get("data", [])
            name = s.get("name", "")
            if s_type == "bar":
                if chart_config.get("xAxis", {}).get("type") == "value":
                    # horizontal bar
                    y_data = chart_config.get("yAxis", {}).get("data", [])
                    ax.barh(y_data, data)
                else:
                    ax.bar(x_data, data, label=name)
            elif s_type == "line":
                ax.plot(x_data, data, label=name, marker="o")
            elif s_type == "pie":
                pie_data = s.get("data", [])
                labels = [d["name"] for d in pie_data]
                values = [d["value"] for d in pie_data]
                ax.pie(values, labels=labels, autopct="%1.1f%%")
                ax.set_title(title)
                plt.tight_layout()
                buf = io.BytesIO()
                fig.savefig(buf, format="png", dpi=150)
                plt.close()
                buf.seek(0)
                return buf.read()

        ax.set_title(title, fontsize=14)
        if chart_config.get("xAxis", {}).get("name"):
            ax.set_xlabel(chart_config["xAxis"]["name"])
        if chart_config.get("yAxis", {}).get("name"):
            ax.set_ylabel(chart_config["yAxis"]["name"])
        if len(series) > 1 or any(s.get("name") for s in series):
            ax.legend()
        plt.xticks(rotation=30, ha="right")
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150)
        plt.close()
        buf.seek(0)
        return buf.read()

    def generate_docx(self, title: str, content: str, metadata: dict = None) -> str:
        """生成Word文档，返回文件路径"""
        filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{title[:20]}.docx"
        filepath = self.output_dir / filename

        doc = Document()
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if metadata:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            meta_run = meta_para.add_run(
                f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            )
            meta_run.font.size = Pt(9)

        doc.add_paragraph()
        for line in content.split("\n"):
            if line.strip().startswith("#"):
                level = min(len(line) - len(line.lstrip("#")), 3)
                doc.add_heading(line.lstrip("#").strip(), level=level)
            else:
                doc.add_paragraph(line)

        doc.save(filepath)
        return str(filepath)

    def generate_pdf(self, title: str, content: str, metadata: dict = None) -> str:
        """生成PDF文档，返回文件路径"""
        filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{title[:20]}.pdf"
        filepath = self.output_dir / filename

        font_name = self._get_cn_font()
        styles = getSampleStyleSheet()
        cn_style = ParagraphStyle(
            "Chinese",
            parent=styles["Normal"],
            fontName=font_name,
            fontSize=11,
            leading=16,
            spaceAfter=6,
        )
        title_style = ParagraphStyle(
            "ChineseTitle",
            parent=styles["Title"],
            fontName=font_name,
            fontSize=18,
            alignment=1,
            spaceAfter=20,
        )
        h1_style = ParagraphStyle(
            "ChineseH1",
            parent=styles["Heading1"],
            fontName=font_name,
            fontSize=16,
            spaceAfter=12,
        )
        h2_style = ParagraphStyle(
            "ChineseH2",
            parent=styles["Heading2"],
            fontName=font_name,
            fontSize=14,
            spaceAfter=10,
        )

        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )
        story = []
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.5 * cm))

        for line in content.split("\n"):
            stripped = line.strip()
            if stripped.startswith("###"):
                story.append(Paragraph(stripped.lstrip("#").strip(), h2_style))
            elif stripped.startswith("##"):
                story.append(Paragraph(stripped.lstrip("#").strip(), h2_style))
            elif stripped.startswith("#"):
                story.append(Paragraph(stripped.lstrip("#").strip(), h1_style))
            elif stripped:
                story.append(Paragraph(stripped, cn_style))
            else:
                story.append(Spacer(1, 0.3 * cm))

        doc.build(story)
        return str(filepath)

    def export(self, narrative: str, title: str, fmt: str = "docx") -> str:
        """统一导出接口"""
        if fmt.lower() == "docx":
            return self.generate_docx(title, narrative)
        elif fmt.lower() == "pdf":
            return self.generate_pdf(title, narrative)
        else:
            raise ValueError(f"不支持的格式: {fmt}")

    def export_report(self, report: dict, fmt: str = "docx") -> str:
        """导出完整报告（含图表和文本）"""
        title = report.get("title", "数据分析报告")
        texts = report.get("texts", {})
        charts = report.get("charts", {})
        data_sources = report.get("data_sources", {})
        generated_at = report.get("generated_at", "")

        # 将图表转为图片
        chart_images = {}
        for chart_id, chart_cfg in charts.items():
            img_bytes = self._chart_to_image(chart_cfg)
            if img_bytes:
                chart_images[chart_id] = img_bytes

        if fmt.lower() == "docx":
            return self._generate_report_docx(title, texts, chart_images, data_sources, generated_at)
        elif fmt.lower() == "pdf":
            return self._generate_report_pdf(title, texts, chart_images, data_sources, generated_at)
        else:
            raise ValueError(f"不支持的格式: {fmt}")

    def _generate_report_docx(self, title, texts, chart_images, data_sources, generated_at):
        filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_报告.docx"
        filepath = self.output_dir / filename

        doc = Document()
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        meta_run = meta.add_run(f"生成时间: {generated_at}")
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)

        # 数据来源
        doc.add_heading("数据概览", level=1)
        for key, val in data_sources.items():
            doc.add_paragraph(f"{key}: {val.get('file', '')} ({val.get('records', 0)} 条记录)")

        from utils.report_layout import interleave_text_with_charts

        # 各章节
        section_order = [
            ("basic", "一、基本统计"),
            ("admission_trend", "二、入院趋势分析"),
            ("patient_features", "三、患者特征分析"),
            ("hospitalization_days", "四、住院天数分析"),
            ("disease_types", "五、疾病类型提取分析"),
            ("readmission", "六、再入院分析"),
            ("discharge", "七、出院情况分析"),
            ("exam", "检查数据分析"),
            ("lab", "检验数据分析"),
            ("summary", "数据质量评估与总结"),
        ]

        for section_key, section_title in section_order:
            text = texts.get(section_key, "")
            # 去掉文本中已有的章节总标题，避免重复
            lines = text.split("\n") if text else []
            filtered_lines = []
            title_found = False
            for line in lines:
                stripped = line.strip()
                # 如果第一行就是章节标题（去除markdown标记后），跳过
                if not title_found and stripped:
                    clean_title = stripped.lstrip("#").strip()
                    if clean_title == section_title or clean_title.replace("一、", "").replace("二、", "").replace("三、", "").replace("四、", "").replace("五、", "").replace("六、", "").replace("七、", "").strip() == section_title.replace("一、", "").replace("二、", "").replace("三、", "").replace("四、", "").replace("五、", "").replace("六、", "").replace("七、", "").strip():
                        title_found = True
                        continue
                    title_found = True
                filtered_lines.append(line)
            text = "\n".join(filtered_lines)

            doc.add_heading(section_title, level=1)
            # 将图表与文本穿插排列
            blocks = interleave_text_with_charts(section_key, text, chart_images)
            for block in blocks:
                if block["type"] == "text":
                    for line in block["content"].split("\n"):
                        if line.strip().startswith("#"):
                            level = min(len(line) - len(line.lstrip("#")), 3)
                            doc.add_heading(line.lstrip("#").strip(), level=level)
                        else:
                            self._add_markdown_paragraph(doc, line)
                elif block["type"] == "chart":
                    ck = block["chart_id"]
                    if ck in chart_images:
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(io.BytesIO(chart_images[ck]), width=Inches(5.5))

        doc.save(filepath)
        return str(filepath)

    def _generate_report_pdf(self, title, texts, chart_images, data_sources, generated_at):
        filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_报告.pdf"
        filepath = self.output_dir / filename

        font_name = self._get_cn_font()
        styles = getSampleStyleSheet()
        cn_style = ParagraphStyle(
            "Chinese", parent=styles["Normal"], fontName=font_name, fontSize=11, leading=16, spaceAfter=6
        )
        title_style = ParagraphStyle(
            "ChineseTitle", parent=styles["Title"], fontName=font_name, fontSize=18, alignment=1, spaceAfter=20
        )
        h1_style = ParagraphStyle(
            "ChineseH1", parent=styles["Heading1"], fontName=font_name, fontSize=16, spaceAfter=12
        )
        h2_style = ParagraphStyle(
            "ChineseH2", parent=styles["Heading2"], fontName=font_name, fontSize=14, spaceAfter=10
        )
        caption_style = ParagraphStyle(
            "Caption", parent=styles["Normal"], fontName=font_name, fontSize=9, textColor=colors.grey, alignment=1
        )

        doc = SimpleDocTemplate(
            str(filepath), pagesize=A4,
            rightMargin=2 * cm, leftMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm,
        )
        story = []
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"生成时间: {generated_at}", caption_style))
        story.append(Spacer(1, 0.5 * cm))

        from utils.report_layout import interleave_text_with_charts

        section_order = [
            ("basic", "一、基本统计"),
            ("admission_trend", "二、入院趋势分析"),
            ("patient_features", "三、患者特征分析"),
            ("hospitalization_days", "四、住院天数分析"),
            ("disease_types", "五、疾病类型提取分析"),
            ("readmission", "六、再入院分析"),
            ("discharge", "七、出院情况分析"),
            ("exam", "检查数据分析"),
            ("lab", "检验数据分析"),
            ("summary", "数据质量评估与总结"),
        ]

        tmp_files = []
        for section_key, section_title in section_order:
            text = texts.get(section_key, "")
            # 去掉文本中已有的章节总标题，避免重复
            lines = text.split("\n") if text else []
            filtered_lines = []
            title_found = False
            for line in lines:
                stripped = line.strip()
                if not title_found and stripped:
                    clean_title = stripped.lstrip("#").strip()
                    if clean_title == section_title or clean_title.replace("一、", "").replace("二、", "").replace("三、", "").replace("四、", "").replace("五、", "").replace("六、", "").replace("七、", "").strip() == section_title.replace("一、", "").replace("二、", "").replace("三、", "").replace("四、", "").replace("五、", "").replace("六、", "").replace("七、", "").strip():
                        title_found = True
                        continue
                    title_found = True
                filtered_lines.append(line)
            text = "\n".join(filtered_lines)

            story.append(Paragraph(section_title, h1_style))
            # 将图表与文本穿插排列
            blocks = interleave_text_with_charts(section_key, text, chart_images)
            for block in blocks:
                if block["type"] == "text":
                    for line in block["content"].split("\n"):
                        stripped = line.strip()
                        if stripped.startswith("###"):
                            story.append(Paragraph(self._markdown_to_html(stripped.lstrip("#").strip()), h2_style))
                        elif stripped.startswith("##"):
                            story.append(Paragraph(self._markdown_to_html(stripped.lstrip("#").strip()), h2_style))
                        elif stripped.startswith("#"):
                            story.append(Paragraph(self._markdown_to_html(stripped.lstrip("#").strip()), h1_style))
                        elif stripped:
                            story.append(Paragraph(self._markdown_to_html(stripped), cn_style))
                        else:
                            story.append(Spacer(1, 0.3 * cm))
                elif block["type"] == "chart":
                    ck = block["chart_id"]
                    if ck in chart_images:
                        img_path = self.output_dir / f"_tmp_{ck}.png"
                        with open(img_path, "wb") as f:
                            f.write(chart_images[ck])
                        tmp_files.append(img_path)
                        story.append(Spacer(1, 0.3 * cm))
                        img = Image(str(img_path), width=14 * cm, height=8 * cm)
                        story.append(img)
                        story.append(Paragraph(f"图: {ck}", caption_style))

            story.append(Spacer(1, 0.5 * cm))

        doc.build(story)

        # 清理临时文件
        for fp in tmp_files:
            try:
                os.remove(fp)
            except Exception:
                pass
        return str(filepath)


# 全局单例
document_service = DocumentService()
