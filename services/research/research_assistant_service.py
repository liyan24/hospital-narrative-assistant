"""
科研助手编排服务：算子执行 + LLM 解读 + 分析路径推荐 + IMRaD 论文生成（docx）。
"""
import io
import uuid
from datetime import datetime
from pathlib import Path

from database.json_store import json_store
from services.llm_service import llm_service
from services.research.dataset_service import dataset_service
from services.research.skills.base import convert
from services.research.skills.registry import get_skill, list_skills_by_category

OUTPUT_DIR = Path("./data/outputs")

SYSTEM_PROMPT = (
    "你是一位资深的临床科研方法学顾问与医学论文写作专家，"
    "熟悉肿瘤与血液病领域、真实世界研究方法和中文医学论文规范（GB/T 7714）。"
    "回答使用中文，严谨专业；只能基于提供的事实数据展开，不得编造数字。"
)


class ResearchAssistantService:
    """科研助手编排服务"""

    def __init__(self):
        self.llm = llm_service

    # ========== 算子执行 + LLM 解读 ==========

    def run_skill(self, skill_id: str, params: dict) -> dict:
        """执行算子，附加 LLM 中文解读，结果存 json_store"""
        skill = get_skill(skill_id)
        if skill is None:
            raise ValueError(f"未知算子: {skill_id}")

        result = skill.run(params or {})

        interpretation = self._interpret(skill.meta, params, result)

        result_id = f"research_result_{uuid.uuid4().hex[:8]}"
        record = {
            "result_id": result_id,
            "skill_id": skill_id,
            "skill_name": skill.meta.name,
            "category": skill.meta.category,
            "params": params or {},
            "result": result,
            "interpretation": interpretation,
            "created_at": datetime.now().isoformat(timespec="seconds"),
        }
        json_store.save(result_id, convert(record))
        return convert(record)

    def _interpret(self, meta, params: dict, result: dict) -> str:
        """LLM 对算子结果生成中文解读（失败返回错误字符串，不中断流程）"""
        # 摘要 + 关键事实 + 首个表格前几行，控制 prompt 规模
        brief_tables = []
        for t in result.get("tables", [])[:2]:
            brief_tables.append({
                "title": t.get("title"),
                "columns": t.get("columns"),
                "rows": t.get("rows", [])[:8],
            })
        import json
        context = json.dumps({
            "summary": result.get("summary", ""),
            "facts": result.get("facts", {}),
            "tables_preview": brief_tables,
        }, ensure_ascii=False, default=str)[:5000]

        return self.llm.chat(
            [{"role": "system", "content": SYSTEM_PROMPT},
             {"role": "user", "content":
              f"以下是科研算子「{meta.name}」（{meta.category}）的执行结果，参数：{params}。\n\n"
              f"{context}\n\n"
              "请用中文给出 200-400 字的专业解读：1) 核心发现；2) 临床/科研意义；"
              "3) 结果解读时的注意事项（混杂、偏倚、多重比较等）；4) 建议的下一步分析。"}],
            temperature=0.4,
            max_tokens=1500,
            cache_namespace="research:interpret",
        )

    def get_result(self, result_id: str) -> dict | None:
        return json_store.load(result_id)

    # ========== 分析路径推荐 ==========

    def recommend_path(self, question: str) -> dict:
        """LLM 根据数据资产 + 算子目录推荐分析路径"""
        assets = dataset_service.detect_data_assets()
        assets_brief = [{
            "name": t["name"], "rows": t["rows"], "coverage_note": t["coverage_note"]
        } for t in assets["tables"]]

        catalog = []
        for category, skills in list_skills_by_category().items():
            for s in skills:
                catalog.append({"id": s["id"], "name": s["name"], "category": category,
                                "description": s["description"]})

        import json
        content = self.llm.chat(
            [{"role": "system", "content": SYSTEM_PROMPT},
             {"role": "user", "content":
              f"研究问题：{question}\n\n"
              f"数据资产：{json.dumps(assets_brief, ensure_ascii=False)}\n"
              f"图谱可用：{assets['graph'].get('available', False)}\n\n"
              f"可用分析算子：\n{json.dumps(catalog, ensure_ascii=False, indent=1)}\n\n"
              "请推荐一条可落地的分析路径：1) 总体思路一段话；"
              "2) 按顺序列出建议执行的算子（使用算子 id），并说明每一步的目的。"
              "只推荐与数据条件匹配的算子（图谱不可用时不要推荐图谱算子）。"}],
            temperature=0.4,
            max_tokens=2000,
            cache_namespace="research:recommend",
        )

        # 从回复中识别提到的算子 id
        suggested = [sid for sid in
                     [s["id"] for skills in list_skills_by_category().values() for s in skills]
                     if sid in content]
        return {
            "recommendation": content,
            "suggested_skills": suggested,
        }

    # ========== 论文生成 ==========

    PAPER_SECTIONS = ["摘要", "前言", "资料与方法", "结果", "讨论", "结论"]

    def generate_paper(self, question: str, result_ids: list[str],
                       articles: list[dict], title: str | None = None) -> dict:
        """汇总选定结果与文献，LLM 分章节生成 IMRaD 论文并导出 docx"""
        import json

        # 汇总选定分析结果
        selected = []
        all_charts = []
        for rid in result_ids:
            record = json_store.load(rid)
            if not record:
                continue
            selected.append({
                "skill": record.get("skill_name"),
                "summary": record.get("result", {}).get("summary", ""),
                "facts": record.get("result", {}).get("facts", {}),
            })
            for ch in record.get("result", {}).get("charts", [])[:3]:
                all_charts.append({"title": ch.get("title", ""), "option": ch.get("option", {})})

        if not selected:
            raise ValueError("未找到任何有效的分析结果，请先运行算子并选择至少一个结果")

        facts_brief = json.dumps(selected, ensure_ascii=False, default=str)[:8000]
        articles_brief = ""
        if articles:
            articles_brief = "\n".join(
                f"[{i+1}] {a.get('title', '')}（{a.get('journal', '')}, {a.get('year', '')}）"
                f"：{(a.get('abstract') or '')[:300]}"
                for i, a in enumerate(articles[:10])
            )

        # 生成题目
        if not title:
            title = self.llm.chat(
                [{"role": "system", "content": SYSTEM_PROMPT},
                 {"role": "user", "content":
                  f"研究问题：{question}\n分析结果：\n{facts_brief[:2000]}\n\n"
                  "请为该研究拟定一个规范的中文医学论文题目（一句话，只输出题目本身）。"}],
                temperature=0.4, max_tokens=200,
                cache_namespace="research:paper_title",
            ).strip().strip("《》\"'")

        context = (
            f"论文题目：{title}\n研究问题：{question}\n\n"
            f"分析结果事实（引用数据必须来自此处，不得编造）：\n{facts_brief}\n\n"
            + (f"参考文献：\n{articles_brief}\n\n" if articles_brief else "")
            + "研究背景：某县级医院肿瘤血液科真实世界回顾性数据"
            "（约1.37万就诊、3990名患者，含诊断/医嘱/检验/检查，无性别字段）。"
        )

        sections = {}
        for sec in self.PAPER_SECTIONS:
            extra = "（含3-5个关键词，300字左右）" if sec == "摘要" else "（500-800字，学术语言）"
            if sec == "结果":
                extra = "（按小节组织，逐一引用上述分析结果中的具体数字；标注'表X'/'图X'占位）"
            content = self.llm.chat(
                [{"role": "system", "content": SYSTEM_PROMPT},
                 {"role": "user", "content":
                  f"{context}\n请撰写论文「{sec}」部分{extra}。只输出该章节正文。"}],
                temperature=0.4, max_tokens=2500,
                cache_namespace="research:paper_writing",
            )
            sections[sec] = content

        # 参考文献（GB/T 7714）
        references = ""
        if articles:
            references = self.llm.chat(
                [{"role": "system", "content": SYSTEM_PROMPT},
                 {"role": "user", "content":
                  "请将以下文献格式化为 GB/T 7714-2015 顺序编码制格式"
                  f"（卷期页缺失可省略，末尾保留 PMID）：\n{json.dumps(articles[:10], ensure_ascii=False)[:5000]}\n"
                  "只输出编号文献列表。"}],
                temperature=0.2, max_tokens=2000,
                cache_namespace="research:reference_format",
            )

        # 生成 docx
        filename = f"科研论文_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = self._build_docx(OUTPUT_DIR / filename, title, sections, references, all_charts[:6])

        project_id = f"research_project_{uuid.uuid4().hex[:8]}"
        json_store.save(project_id, convert({
            "project_id": project_id,
            "question": question,
            "title": title,
            "result_ids": result_ids,
            "sections": sections,
            "references": references,
            "filename": filename,
            "created_at": datetime.now().isoformat(timespec="seconds"),
        }))

        return convert({
            "paper": {"title": title, "sections": sections, "references": references},
            "filename": filename,
            "download_url": f"/api/research/paper/download/{filename}",
        })

    # ========== docx 生成（参照 document_service 的排版与 _chart_to_image 模式） ==========

    def _build_docx(self, filepath: Path, title: str, sections: dict,
                    references: str, charts: list) -> str:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt

        filepath.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()

        def set_font(run, cn_font: str, size: int, bold: bool = False):
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
            run.font.size = Pt(size)
            run.font.bold = bold

        # 标题（黑体二号居中）
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(title), "黑体", 16, bold=True)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}（AI 辅助起草，需人工审校）"),
                 "宋体", 9)

        def add_section(sec_title: str, content: str):
            h = doc.add_paragraph()
            set_font(h.add_run(sec_title), "黑体", 14, bold=True)
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue
                para = doc.add_paragraph()
                para.paragraph_format.first_line_indent = Pt(24)
                para.paragraph_format.line_spacing = 1.5
                set_font(para.add_run(line), "宋体", 12)

        for sec in self.PAPER_SECTIONS:
            content = sections.get(sec, "")
            if not content or content.startswith("[LLM调用失败]"):
                content = f"（本章节生成失败：{content or '空'}，请补充。）"
            add_section(sec, content)
            # 结果章节后嵌入图表
            if sec == "结果" and charts:
                for ch in charts:
                    img = self._chart_to_image(ch.get("option", {}))
                    if img:
                        doc.add_picture(io.BytesIO(img), width=Pt(400))
                        cap = doc.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        set_font(cap.add_run(ch.get("title", "图")), "黑体", 10)

        if references:
            add_section("参考文献", references)

        doc.save(str(filepath))
        return str(filepath)

    def _chart_to_image(self, option: dict, width: int = 8, height: int = 5) -> bytes:
        """将 ECharts option 用 matplotlib 渲染为 PNG 字节（参照 document_service._chart_to_image）"""
        import matplotlib
        matplotlib.use("Agg")
        matplotlib.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "Arial Unicode MS", "DejaVu Sans"]
        matplotlib.rcParams["axes.unicode_minus"] = False
        import matplotlib.pyplot as plt

        series = option.get("series", [])
        if not series:
            return b""

        title = option.get("title", {}).get("text", "")
        x_data = option.get("xAxis", {}).get("data", [])
        s0 = series[0]
        s_type = s0.get("type", "bar")

        fig, ax = plt.subplots(figsize=(width, height))
        try:
            if s_type == "pie":
                data = s0.get("data", [])
                ax.pie([d["value"] for d in data], labels=[d["name"] for d in data],
                       autopct="%1.1f%%")
                ax.set_title(title)
            elif s_type == "scatter":
                for s in series:
                    data = s.get("data", [])
                    if not data:
                        continue
                    xs = [p[0] for p in data]
                    ys = [p[1] for p in data]
                    ax.scatter(xs, ys, s=8, alpha=0.5, label=s.get("name", ""))
                ax.set_title(title)
                ax.set_xlabel(option.get("xAxis", {}).get("name", ""))
                ax.set_ylabel(option.get("yAxis", {}).get("name", ""))
                if any(s.get("name") for s in series):
                    ax.legend()
            elif s_type == "heatmap":
                data = s0.get("data", [])
                x_labels = option.get("xAxis", {}).get("data", [])
                y_labels = option.get("yAxis", {}).get("data", [])
                import numpy as np
                mat = np.zeros((len(y_labels), len(x_labels)))
                for xi, yi, v in data:
                    mat[yi][xi] = v
                im = ax.imshow(mat, cmap="RdBu_r", vmin=-1, vmax=1)
                ax.set_xticks(range(len(x_labels)), x_labels, rotation=45, ha="right")
                ax.set_yticks(range(len(y_labels)), y_labels)
                fig.colorbar(im, ax=ax)
                ax.set_title(title)
            elif s_type == "graph":
                # 网络图不适合静态渲染，跳过
                plt.close()
                return b""
            elif s_type == "bar" and option.get("xAxis", {}).get("type") == "value":
                y_data = option.get("yAxis", {}).get("data", [])
                ax.barh([str(y) for y in y_data], s0.get("data", []))
                ax.set_title(title)
            else:
                for s in series:
                    data = s.get("data", [])
                    if s.get("type") == "line":
                        ax.plot(x_data, data, label=s.get("name", ""), marker="o")
                    else:
                        ax.bar([str(x) for x in x_data], data, label=s.get("name", ""))
                ax.set_title(title)
                ax.set_xlabel(option.get("xAxis", {}).get("name", ""))
                ax.set_ylabel(option.get("yAxis", {}).get("name", ""))
                if len(series) > 1 or any(s.get("name") for s in series):
                    ax.legend()
                plt.xticks(rotation=30, ha="right")

            plt.tight_layout()
            buf = io.BytesIO()
            fig.savefig(buf, format="png", dpi=150)
            plt.close()
            buf.seek(0)
            return buf.read()
        except Exception:
            plt.close()
            return b""


# 全局单例
research_assistant_service = ResearchAssistantService()
