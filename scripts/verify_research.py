"""科研助手模块自验证脚本（一次性验证用，可删除）。"""
import json
import sys
import time
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

results = {}

# ① main:app 可 import
from main import app  # noqa
routes = [r.path for r in app.routes if "/api/research" in getattr(r, "path", "")]
results["1_app_import"] = {"ok": True, "research_routes": len(routes)}
print(f"[1] main:app import OK, 科研路由数={len(routes)}")

# ② registry 24 个 skill 全部注册
from services.research.skills.registry import SKILL_REGISTRY, list_skills_by_category
results["2_registry"] = {"count": len(SKILL_REGISTRY), "ids": sorted(SKILL_REGISTRY.keys())}
print(f"[2] registry skill 数={len(SKILL_REGISTRY)}")
assert len(SKILL_REGISTRY) == 24, f"期望24个算子，实际{len(SKILL_REGISTRY)}"

# ③ build_visit_matrix 行数 ≈ 13743
from services.research.dataset_service import dataset_service
t0 = time.time()
visits = dataset_service.build_visit_matrix()
t1 = time.time()
results["3_visit_matrix"] = {
    "rows": len(visits), "cols": list(visits.columns),
    "first_build_seconds": round(t1 - t0, 1),
    "age_median": round(float(visits["age_years"].median()), 1),
    "readmission_rate": round(float(visits["is_readmission"].mean()), 3),
    "drug_visits": int((visits["drugs"].apply(len) > 0).sum()),
    "lab_cols": [c for c in visits.columns if c.startswith("lab_")],
}
print(f"[3] visit_matrix 行数={len(visits)} 列={list(visits.columns)}")
assert abs(len(visits) - 13743) < 100

# 缓存第二次调用秒回
t0 = time.time()
_ = dataset_service.build_visit_matrix()
_ = dataset_service.load_table("orders")
results["3_cache_second_call_seconds"] = round(time.time() - t0, 2)
print(f"[3b] 缓存第二次调用耗时={results['3_cache_second_call_seconds']}s")

# ④ 四个算子直接调用返回合法结构
from services.research.skills.registry import get_skill
for sid in ["dataset_profile", "association_rules", "classification", "clustering"]:
    t0 = time.time()
    out = get_skill(sid).run({})
    ok = all(k in out for k in ("summary", "tables", "charts", "facts"))
    results[f"4_{sid}"] = {
        "ok": ok, "seconds": round(time.time() - t0, 1),
        "summary_head": out["summary"][:120],
        "tables": len(out["tables"]), "charts": len(out["charts"]),
    }
    print(f"[4] {sid}: ok={ok} tables={len(out['tables'])} charts={len(out['charts'])} "
          f"summary={out['summary'][:80]}...")
    assert ok

# ⑤ custom_code 拒绝 import os，且能执行正常 pandas 代码
from services.research.custom_code_service import custom_code_service
bad = custom_code_service.run("import os\nresult = os.listdir()")
results["5_reject_import"] = {"summary": bad["summary"][:100]}
print(f"[5] 拒绝 import os: {bad['summary'][:80]}")
assert "拒绝" in bad["summary"] or "禁止" in bad["summary"]

good = custom_code_service.run(
    "df = visits.groupby('is_readmission')['length_of_stay'].median().reset_index()\nresult = df")
results["5_run_pandas"] = {
    "summary": good["summary"][:100],
    "table_cols": good["tables"][0]["columns"] if good["tables"] else None,
}
print(f"[5] 正常 pandas: {good['summary'][:80]}")
assert good["tables"], "pandas 代码未返回表格"

# ⑥ run_skill 全链路（含 LLM 真实解读）
from services.research.research_assistant_service import research_assistant_service
t0 = time.time()
full = research_assistant_service.run_skill("descriptive_stats", {})
interp = full["interpretation"]
results["6_run_skill"] = {
    "result_id": full["result_id"],
    "seconds": round(time.time() - t0, 1),
    "interp_head": interp[:150],
    "interp_failed": interp.startswith("[LLM调用失败]"),
}
print(f"[6] run_skill result_id={full['result_id']} 解读失败={results['6_run_skill']['interp_failed']}")
print(f"    解读开头: {interp[:120]}")
assert not interp.startswith("[LLM调用失败]"), "LLM 解读失败"

# ⑦ generate_paper 产出 docx 并读回验证章节
t0 = time.time()
paper = research_assistant_service.generate_paper(
    question="肿瘤血液科患者再入院的影响因素分析",
    result_ids=[full["result_id"]],
    articles=[],
    title=None,
)
results["7_paper"] = {
    "filename": paper["filename"], "title": paper["paper"]["title"],
    "seconds": round(time.time() - t0, 1),
    "sections": list(paper["paper"]["sections"].keys()),
}
print(f"[7] 论文: {paper['paper']['title']} -> {paper['filename']}")

from docx import Document
doc = Document(f"./data/outputs/{paper['filename']}")
text = "\n".join(p.text for p in doc.paragraphs)
found = [s for s in ["摘要", "前言", "资料与方法", "结果", "讨论", "结论"] if s in text]
results["7_docx_sections_found"] = found
print(f"[7] docx 章节验证: {found}, 段落数={len(doc.paragraphs)}")
assert len(found) == 6, f"docx 章节不全: {found}"

print("\n===== 验证全部通过 =====")
with open("data/outputs/research_selfcheck.json", "w", encoding="utf-8") as f:
    json.dump(results, f, ensure_ascii=False, indent=2)
print("结果已写入 data/outputs/research_selfcheck.json")
